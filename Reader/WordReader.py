from Document import Document
from docx import Document as DocxDocument
import os

class WordReader:
    def __init__(self, file_paths):
        self.file_paths = file_paths

    def read_word_document(self, file_path):
        # Đọc tài liệu Word
        doc = DocxDocument(file_path)
        paragraphs = []

        for para in doc.paragraphs:
            # Thêm đối tượng đoạn văn vào danh sách
            paragraphs.append(para)

        return paragraphs

    def is_bold_and_uppercase(self, paragraph):
        # Kiểm tra xem đoạn văn có chứa văn bản in đậm và viết hoa toàn bộ không
        text = paragraph.text.strip()
        return any(run.bold for run in paragraph.runs) and text.isupper()

    def split_by_format(self, paragraphs):
        # Phân đoạn văn bản dựa trên định dạng
        sections = []
        current_section = []

        for paragraph in paragraphs:
            paragraph_text = paragraph.text
            if self.is_bold_and_uppercase(paragraph):
                if current_section:
                    sections.append("\n".join([p.text for p in current_section]))
                    current_section = []
                current_section.append(paragraph)
            else:
                current_section.append(paragraph)

        if current_section:
            sections.append("\n".join([p.text for p in current_section]))

        return sections

    def process_documents(self):
        all_sections = []

        for file_path in self.file_paths:
            if not os.path.exists(file_path):
                print(f"File does not exist: {file_path}")
                continue

            paragraphs = self.read_word_document(file_path)
            sections = self.split_by_format(paragraphs)
            # Tạo các đối tượng Document từ các đoạn văn bản
            documents = [Document(section) for section in sections]
            all_sections.extend(documents)

        return all_sections
